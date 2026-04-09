# pages/statistiques_prescripteur_page.py - VERSION DÉFINITIVE CORRIGÉE
import streamlit as st
import pandas as pd
from io import BytesIO
from modules.data_manager import load_data
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

def get_current_decade_info():
    """
    Détermine la décade actuelle en fonction de la date du jour
    Retourne un tuple (periode, date_debut, date_fin)
    """
    data = load_data()
    today = datetime.now().strftime("%Y-%m-%d")
    
    for decade in data["decades"]:
        if decade["date_debut"] <= today <= decade["date_fin"]:
            return {
                "periode": decade["periode"],
                "date_debut": decade["date_debut"],
                "date_fin": decade["date_fin"],
                "mois": decade["mois"],
                "annee": decade["annee"]
            }
    
    # Si aucune décade trouvée, retourner None
    return None

def get_current_month_and_year():
    """
    Retourne le mois et l'année de la date d'impression
    """
    now = datetime.now()
    mois_fr = {1:"Janvier", 2:"Février", 3:"Mars", 4:"Avril", 5:"Mai", 6:"Juin", 
               7:"Juillet", 8:"Août", 9:"Septembre", 10:"Octobre", 11:"Novembre", 12:"Décembre"}
    return mois_fr[now.month], now.year

def get_decade_label_with_interval(decade_name):
    """
    Retourne le libellé de la décade avec son intervalle
    """
    if "1ère" in decade_name or "1ere" in decade_name:
        return "1ère décade\n(1 - 10)"
    elif "2ème" in decade_name or "2eme" in decade_name:
        return "2ème décade\n(11 - 20)"
    elif "3ème" in decade_name or "3eme" in decade_name:
        return "3ème décade\n(21 - 30/31)"
    return decade_name

def show():
    # --- CSS SPÉCIFIQUE STATS ---
    st.markdown("""
    <style>
        .card-stat {
            background: white;
            padding: 20px;
            border-radius: 16px;
            border: 1px solid #e2e8f0;
            box-shadow: 0 4px 6px -1px rgba(0,0,0,0.05);
        }
        .stat-label { font-size: 0.75rem; color: #64748b; font-weight: 700; text-transform: uppercase; }
        .stat-val { font-size: 1.5rem; font-weight: 800; color: #0f172a; display: block; margin-top: 5px; }
        .export-card {
            background: #f8fafc;
            padding: 20px;
            border-radius: 16px;
            border: 1px dashed #cbd5e1;
            text-align: center;
        }
    </style>
    """, unsafe_allow_html=True)

    c1, c2 = st.columns([3, 1])
    with c1:
        st.title("👨‍⚕️ Stats Prescripteurs")
        st.caption("Performance et activité détaillée par prescripteur")
    with c2:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🔄 Actualiser", use_container_width=True):
            st.rerun()

    data = load_data()
    if not data["transactions"]:
        st.info("ℹ️ Aucune transaction enregistrée.")
        return

    # Chargement des données
    examens = {e["id"]: e for e in data["examens"]}
    prescripteurs = {p["id"]: p["nom"] for p in data["prescripteurs"]}
    decades = {d["id"]: d for d in data["decades"]}
    categories = {c["id"]: c["nom"] for c in data["categories"]}

    records = []
    for t in data["transactions"]:
        ex = examens.get(t["examen_id"])
        if not ex: continue
        dec = decades.get(t["decade_id"], {})
        pres_nom = prescripteurs.get(t["prescripteur_id"], "Inconnu")
        cat_nom = categories.get(ex.get("categorie_id"), "Autres")
        montant = t["quantite_payee"] * ex.get("prix", 0)

        records.append({
            "prescripteur": pres_nom,
            "examen": ex.get("nom", ""),
            "categorie": cat_nom,
            "decade": dec.get("periode", ""),
            "mois": dec.get("mois", ""),
            "annee": dec.get("annee", ""),
            "qte_payee": t["quantite_payee"],
            "qte_gratuite": t["quantite_gratuite"],
            "montant": montant
        })

    df_raw = pd.DataFrame(records)
    if df_raw.empty:
        st.info("Aucune transaction trouvée.")
        return

    # --- FILTRES MODERNES ---
    with st.expander("🔍 Options de filtrage", expanded=True):
        f1, f2, f3 = st.columns(3)
        with f1:
            prescripteurs_list = sorted(df_raw["prescripteur"].unique())
            prescripteur_sel = st.selectbox("👤 Prescripteur", ["Tous les prescripteurs"] + prescripteurs_list)
        with f2:
            annee_list = sorted(df_raw["annee"].unique(), reverse=True)
            annee_sel = st.selectbox("📆 Année", ["Toutes"] + annee_list)
        with f3:
            mois_list = sorted(df_raw["mois"].unique())
            mois_sel = st.selectbox("📅 Mois", ["Tous"] + mois_list)

    df_filtered = df_raw.copy()
    if prescripteur_sel != "Tous les prescripteurs":
        df_filtered = df_filtered[df_filtered["prescripteur"] == prescripteur_sel]
    if annee_sel != "Toutes":
        df_filtered = df_filtered[df_filtered["annee"] == int(annee_sel)]
    if mois_sel != "Tous":
        df_filtered = df_filtered[df_filtered["mois"] == mois_sel]

    if df_filtered.empty:
        st.warning("Aucune donnée pour ces filtres.")
        return

    # --- KPI CARDS ---
    total_montant = df_filtered["montant"].sum()
    total_qte_payee = df_filtered["qte_payee"].sum()
    total_qte_gratuite = df_filtered["qte_gratuite"].sum()
    nb_prescripteurs = df_filtered["prescripteur"].nunique()

    k1, k2, k3, k4 = st.columns(4)
    with k1:
        st.markdown(f'<div class="card-stat"><span class="stat-label">Actifs</span><span class="stat-val">{nb_prescripteurs}</span></div>', unsafe_allow_html=True)
    with k2:
        st.markdown(f'<div class="card-stat"><span class="stat-label">Payés</span><span class="stat-val">{int(total_qte_payee)}</span></div>', unsafe_allow_html=True)
    with k3:
        st.markdown(f'<div class="card-stat"><span class="stat-label">Gratuits</span><span class="stat-val">{int(total_qte_gratuite)}</span></div>', unsafe_allow_html=True)
    with k4:
        st.markdown(f'<div class="card-stat"><span class="stat-label">Total FCFA</span><span class="stat-val" style="color:#2563eb;">{total_montant:,.0f}</span></div>', unsafe_allow_html=True)

    # --- AFFICHAGE ET EXPORT WORD UNIQUEMENT ---
    st.markdown("<br>", unsafe_allow_html=True)
    col_tab, col_exp = st.columns([2, 1])
    
    with col_tab:
        st.subheader("📋 Récapitulatif par période")
        pivot_display = df_filtered.groupby(["prescripteur", "decade"])["montant"].sum().unstack(fill_value=0)
        st.dataframe(pivot_display, use_container_width=True)

    with col_exp:
        st.subheader("📎 Export")
        st.markdown('<div class="export-card">', unsafe_allow_html=True)
        st.write("Rapport par prescripteur (Modèle institutionnel).")
        st.markdown("<br>", unsafe_allow_html=True)
        
        if st.button("📄 Télécharger WORD", type="primary", key="btn_wd_pres", use_container_width=True):
            export_prescripteur_word(df_filtered)
            st.toast("✅ Rapport Word généré avec succès", icon="📝")
        st.markdown('</div>', unsafe_allow_html=True)


# ====================== EXPORT WORD - VERSION AVEC INTERVALLES ET DÉCADE ACTUELLE ======================
def export_prescripteur_word(df_source):
    """Génère le rapport Word avec intervalles des décades et décade actuelle dans le titre"""
    if df_source.empty:
        st.warning("Aucune donnée à exporter.")
        return

    # --- 1. PRÉPARATION ---
    df_work = df_source.copy()
    mask_kit = df_work['examen'].str.contains('Kit|Prélèvement|PRELEVEMENT', case=False, na=False)
    df_kits = df_work[mask_kit]
    df_presc = df_work[~mask_kit].copy()

    df_presc['categorie'] = df_presc['prescripteur'].apply(
        lambda x: "INFIRMIERS" if str(x).startswith("Mme") else "MEDECINS"
    )
    
    dec_order = ["1ère décade", "2ème décade", "3ème décade"]
    dec_labels_with_interval = ["1ère décade (1-10)", "2ème décade (11-20)", "3ème décade (21-30/31)"]
    
    # Récupérer le mois et l'année de la date d'impression
    mois_impression, annee_impression = get_current_month_and_year()
    
    # Récupérer la décade actuelle
    current_decade_info = get_current_decade_info()
    current_decade_text = ""
    if current_decade_info:
        current_decade_text = f" - {current_decade_info['periode']}"
    
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.4)

    # --- 2. EN-TÊTE ADMINISTRATIF ---
    header_table = doc.add_table(rows=1, cols=3)
    header_table.width = Inches(10)
    
    p_fr = header_table.cell(0, 0).paragraphs[0]
    p_fr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    txt_fr = ("REPUBLIQUE DU CAMEROUN\nPaix – Travail – Patrie\n------------\n"
              "MINISTERE DE LA SANTE PUBLIQUE\n------------\nSECRETARIAT GENERAL\n------------\n"
              "DELEGATION REGIONALE DE LA SANTE\nPUBLIQUE DU CENTRE\n------------\n"
              "DISTRICT DE SANTE D'ODZA\n------------\n"
              "CENTRE MEDICAL D'ARRONDISSEMENT DE NKOMO\nTEL. : 671742644")
    run_fr = p_fr.add_run(txt_fr)
    run_fr.bold = True; run_fr.font.size = Pt(7.5)

    if os.path.exists("logo.png"):
        logo_para = header_table.cell(0, 1).paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.add_run().add_picture("logo.png", width=Inches(1.2))

    p_en = header_table.cell(0, 2).paragraphs[0]
    p_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    txt_en = ("REPUBLIC OF CAMEROON\nPeace-Work-Fatherland\n------------\n"
              "MINISTRY OF PUBLIC HEALTH\n------------\nSECRETARIAT GENERAL\n------------\n"
              "CENTER REGIONAL DELEGATION OF\nPUBLIC HEALTH\n------------\n"
              "ODZA HEALTH DISTRICT\n------------\n"
              "NKOMO SUB-DIVISIONAL MEDICAL CENTER")
    run_en = p_en.add_run(txt_en)
    run_en.bold = True; run_en.font.size = Pt(7.5)

    # --- 3. TITRE ET MOIS AVEC DÉCADE ACTUELLE ---
    doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_para.add_run("COMPTE RENDU FINANCIER DES PRESCRIPTEURS AU SERVICE DU LABORATOIRE DU CMA DE NKOMO")
    run_title.bold = True
    run_title.underline = True
    run_title.font.size = Pt(13)

    # Ajouter le mois et l'année de la date d'impression avec la décade actuelle
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = date_para.add_run(f"(MOIS DE {mois_impression.upper()} {annee_impression}{current_decade_text})")
    run_date.bold = True
    run_date.font.size = Pt(11)

    # --- 4. CALCULS COMBINÉS ---
    stats_decades = {dec: {"paye": 0, "gratuit": 0} for dec in dec_order}
    for dec in dec_order:
        d_df = df_presc[df_presc['decade'] == dec]
        stats_decades[dec]["paye"] = int(d_df['qte_payee'].sum())
        stats_decades[dec]["gratuit"] = int(d_df['qte_gratuite'].sum())

    # --- 5. TABLEAU : NOMBRE D'EXAMENS EFFECTUES AVEC INTERVALLES ---
    doc.add_paragraph()
    title_ex = doc.add_paragraph("NOMBRE D'EXAMENS EFFECTUES")
    title_ex.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_ex.runs[0].bold = True; title_ex.runs[0].underline = True

    exam_table = doc.add_table(rows=2, cols=12)
    exam_table.style = 'Table Grid'
    
    # En-têtes décades avec intervalles
    eh1 = exam_table.rows[0].cells
    for i, d_lbl in enumerate(dec_labels_with_interval + ["TOTAL MENSUEL"]):
        eh1[i*3].merge(eh1[i*3+2]).text = d_lbl
        eh1[i*3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_labels = ["Payé", "Gratuit", "TOTAL"] * 4
    for i, lbl in enumerate(sub_labels):
        cell_p = exam_table.rows[1].cells[i].paragraphs[0]
        cell_p.add_run(lbl).bold = True
        cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Remplissage des données
    row_val = exam_table.add_row().cells
    g_p, g_g = 0, 0
    for i, dec in enumerate(dec_order):
        p = stats_decades[dec]["paye"]
        g = stats_decades[dec]["gratuit"]
        row_val[i*3].text, row_val[i*3+1].text, row_val[i*3+2].text = str(p), str(g), str(p+g)
        g_p += p; g_g += g
    row_val[9].text, row_val[10].text, row_val[11].text = str(g_p), str(g_g), str(g_p+g_g)
    for c in row_val: 
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if c.text: c.paragraphs[0].runs[0].bold = True

    # --- 6. SECTION FINANCIÈRE DÉSAGRÉGÉE AVEC INTERVALLES ---
    doc.add_paragraph()
    trans = doc.add_paragraph("RAPPORT FINANCIER DESAGREGE PAR CATEGORIE ET NOM DES PRESCRIPTEURS")
    trans.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trans.runs[0].bold = True; trans.runs[0].underline = True
    doc.add_paragraph("(Nombre d'examens effectues et montants)").alignment = WD_ALIGN_PARAGRAPH.CENTER

    tot_med = {i: 0 for i in range(1, 13)}
    tot_inf = {i: 0 for i in range(1, 13)}
    tot_kit = {i: 0 for i in range(1, 13)}

    categories_to_show = [
        ("MEDECINS", "CATEGORIE DE PRESCRIPTEURS - MEDECINS"),
        ("INFIRMIERS", "CATEGORIE DE PRESCRIPTEURS - INFIRMIERS (SELON LES SERVICES)")
    ]

    for idx, (cat_id, cat_lbl) in enumerate(categories_to_show):
        if idx > 0: doc.add_paragraph()
        
        p_cat = doc.add_paragraph()
        run_cat = p_cat.add_run(cat_lbl)
        run_cat.bold = True; run_cat.underline = True
        
        table = doc.add_table(rows=2, cols=13)
        table.style = 'Table Grid'
        
        h1 = table.rows[0].cells
        h1[0].text = "NOM ET PRENOMS"
        for i, d_lbl in enumerate(dec_labels_with_interval + ["TOTAL MENSUEL"]):
            h1[1+(i*3)].merge(h1[3+(i*3)]).text = d_lbl
        
        h2 = table.rows[1].cells
        labels = ["Payé", "Gratuit", "Montant"] * 4
        for i, txt in enumerate(labels):
            h2[i+1].paragraphs[0].add_run(txt).bold = True

        df_cat = df_presc[df_presc['categorie'] == cat_id]
        if not df_cat.empty:
            summ = df_cat.groupby(["prescripteur", "decade"]).agg({"qte_payee":"sum", "qte_gratuite":"sum", "montant":"sum"}).reset_index()
            for presc in summ["prescripteur"].unique():
                r = table.add_row().cells
                r[0].text = str(presc)
                lp, lg, lm = 0, 0, 0
                for i, dec in enumerate(dec_order):
                    d = summ[(summ["prescripteur"]==presc) & (summ["decade"]==dec)]
                    p, g, m = int(d["qte_payee"].sum()), int(d["qte_gratuite"].sum()), int(d["montant"].sum())
                    r[1+(i*3)].text = str(p) if p>0 else ""
                    r[2+(i*3)].text = str(g) if g>0 else ""
                    r[3+(i*3)].text = f"{m:,}" if m>0 else ""
                    target = tot_med if cat_id == "MEDECINS" else tot_inf
                    target[1+(i*3)]+=p; target[2+(i*3)]+=g; target[3+(i*3)]+=m
                    lp+=p; lg+=g; lm+=m
                r[10].text, r[11].text, r[12].text = str(lp), str(lg), f"{lm:,}"
                target[10]+=lp; target[11]+=lg; target[12]+=lm

        # Ligne Sous-Total
        sr = table.add_row().cells
        sr[0].paragraphs[0].add_run(f"Sous-Total - {cat_id}").bold = True
        curr = tot_med if cat_id == "MEDECINS" else tot_inf
        for i in range(1, 13):
            val_txt = f"{curr[i]:,}" if i%3==0 else str(curr[i])
            sr[i].paragraphs[0].add_run(val_txt).bold = True

        if cat_id == "INFIRMIERS":
            for i, dec in enumerate(dec_order):
                dk = df_kits[df_kits['decade'] == dec]
                pk, gk, mk = int(dk['qte_payee'].sum()), int(dk['qte_gratuite'].sum()), int(dk['montant'].sum())
                tot_kit[1+(i*3)], tot_kit[2+(i*3)], tot_kit[3+(i*3)] = pk, gk, mk
                tot_kit[10]+=pk; tot_kit[11]+=gk; tot_kit[12]+=mk

            for lbl, vals in [("TOTAL GENERAL", {i: tot_med[i]+tot_inf[i] for i in range(1,13)}),
                              ("KIT PRELEVEMENT", tot_kit),
                              ("TOTAL GLOBAL", {i: tot_med[i]+tot_inf[i]+tot_kit[i] for i in range(1,13)})]:
                r = table.add_row().cells
                r[0].paragraphs[0].add_run(lbl).bold = True
                for i in range(1, 13):
                    v_txt = f"{vals[i]:,}" if i%3==0 else str(vals[i])
                    r[i].paragraphs[0].add_run(v_txt).bold = True

    # --- 7. SORTIE ---
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    
    # Générer le nom du fichier avec la date d'impression
    date_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"RAPPORT_PRESCRIPTEURS_{date_str}.docx"
    
    # Téléchargement direct
    st.download_button(
        label="⬇️ Télécharger le Rapport Word Officiel",
        data=out,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )