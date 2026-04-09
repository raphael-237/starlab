import streamlit as st
import pandas as pd
from modules.data_manager import load_data
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

def calculate_stats(data):
    transactions = data.get("transactions", [])
    examens = {e["id"]: e for e in data.get("examens", [])}
    categories = {c["id"]: c["nom"] for c in data.get("categories", [])}
    
    if not transactions: return None

    filtered_data = []
    dates_obj = []
    
    for t in transactions:
        exam = examens.get(t["examen_id"], {})
        cat_nom = categories.get(exam.get("categorie_id"), "").strip().lower()
        if "kit" in cat_nom: continue
            
        filtered_data.append(t)
        if "date" in t:
            try: dates_obj.append(datetime.strptime(t["date"], "%Y-%m-%d"))
            except: pass

    if not filtered_data: return None

    # Dates dynamiques basées sur la BD
    date_debut = min(dates_obj).strftime("%d/%m/%Y") if dates_obj else "01/01/2025"
    date_fin = max(dates_obj).strftime("%d/%m/%Y") if dates_obj else "31/12/2025"

    total_ex = sum(t["quantite_payee"] for t in filtered_data)
    
    cat_counts = {}
    for t in filtered_data:
        c_nom = categories.get(examens.get(t["examen_id"], {}).get("categorie_id"), "Autres")
        cat_counts[c_nom] = cat_counts.get(c_nom, 0) + t["quantite_payee"]

    table_rows = []
    for name, val in cat_counts.items():
        table_rows.append({"nom": name, "total": val, "pc": (val/total_ex)*100})
    
    table_rows.sort(key=lambda x: x["total"], reverse=True)

    conclusion = ""
    if len(table_rows) >= 2:
        poids = table_rows[0]['pc'] + table_rows[1]['pc']
        conclusion = f"La {table_rows[0]['nom'].lower()} et la {table_rows[1]['nom'].lower()} représentent à elles seules près de {int(poids)} % de l'activité annuelle."

    return {
        "total": total_ex, 
        "table": table_rows, 
        "conclusion": conclusion,
        "debut": date_debut,
        "fin": date_fin
    }

def generate_docx(stats):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.different_first_page_header_footer = True
    
    # --- ON NE TOUCHE PAS A L'ENTETE EXISTANTE ---
    header = section.first_page_header
    htable = header.add_table(1, 3, width=Inches(9.0))
    htable.allow_autofit = False
    htable.columns[0].width = Inches(3.7)
    htable.columns[1].width = Inches(1.6)
    htable.columns[2].width = Inches(3.7)

    c_fr = htable.cell(0, 0).paragraphs[0]
    c_fr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_fr.add_run("REPUBLIQUE DU CAMEROUN\nPaix – Travail – Patrie\n------------\nMINISTERE DE LA SANTE PUBLIQUE\n------------\nSECRETARIAT GENERAL\n------------\nDELEGATION REGIONALE DE LA SANTE\nPUBLIQUE DU CENTRE\n------------\nDISTRICT DE SANTE D’ODZA\n------------\nCENTRE MEDICAL D’ARRONDISSEMENT DE NKOMO\nTEL. : 671742644").font.size = Pt(8)

    try:
        c_logo = htable.cell(0, 1).paragraphs[0]
        c_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_logo.add_run().add_picture("logo.png", width=Inches(1.2))
    except: pass

    c_en = htable.cell(0, 2).paragraphs[0]
    c_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c_en.add_run("REPUBLIC OF CAMEROON\nPeace-Work-Fatherland\n------------\nMINISTRY OF PUBLIC HEALTH\n------------\nSECRETARIAT GENERAL\n------------\nCENTER REGIONAL DELEGATION OF\nPUBLIC HEALTH\n------------\nODZA HEALTH DISTRICT\n------------\nNKOMO SUB-DIVISIONAL MEDICAL CENTER").font.size = Pt(8)

    # --- TITRE EXACT (IMAGE) ---
    doc.add_paragraph("\n")
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = titre.add_run("RAPPORT ANNUEL DES EXAMENS RÉALISÉS AU NIVEAU DU LABORATOIRE DU CMA DE NKOMO")
    run_t.bold = True
    run_t.underline = True
    run_t.font.size = Pt(14)

    # --- PERIODE DYNAMIQUE ---
    sous_titre = doc.add_paragraph()
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_st = sous_titre.add_run(f"DU {stats['debut']} AU {stats['fin']}")
    run_st.bold = True
    run_st.font.size = Pt(12)

    # --- 1. ANALYSE ANNUELLE GLOBALE ---
    doc.add_paragraph("\n1. Analyse annuelle globale", style='Heading 1')
    
    # Volume Total (Puce principale)
    p_vol = doc.add_paragraph(style='List Bullet')
    p_vol.add_run("Volume total").bold = True
    
    # Détails (Tirets alignés comme l'image)
    p_a = doc.add_paragraph(f"Total général annuel : {stats['total']:,} examens", style='Normal')
    p_a.paragraph_format.left_indent = Inches(0.5)
    p_a.add_run("").add_break() # Simule le tiret si style absent
    
    p_b = doc.add_paragraph(f"Moyenne mensuelle : {stats['total']/12:,.0f} examens", style='Normal')
    p_b.paragraph_format.left_indent = Inches(0.5)
    
    p_c = doc.add_paragraph(f"Moyenne journalière estimée (base 365 jours ouvrables) : = {stats['total']/365:,.1f} examens/jour", style='Normal')
    p_c.paragraph_format.left_indent = Inches(0.5)

    # Répartition (Puce principale)
    p_rep = doc.add_paragraph(style='List Bullet')
    p_rep.add_run("Répartition par catégorie d'examens").bold = True

    # --- TABLEAU ---
    tab = doc.add_table(rows=1, cols=3)
    tab.style = 'Table Grid'
    hdr = tab.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Catégorie d'examens", "Total annuel", "Part (%)"
    
    for r in stats["table"]:
        row = tab.add_row().cells
        row[0].text = r["nom"]
        row[1].text = f"{r['total']:,}"
        row[2].text = f"{r['pc']:.1f} %"

    f_row = tab.add_row().cells
    f_row[0].text, f_row[1].text, f_row[2].text = "Total", f"{stats['total']:,}", "100 %"
    for cell in f_row:
        for p in cell.paragraphs:
            for run in p.runs: run.bold = True

    # --- CONCLUSION ---
    doc.add_paragraph("\n")
    p_conc = doc.add_paragraph()
    p_conc.add_run(stats["conclusion"]).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def show():
    st.header("📊 Générateur de Rapport")
    data = load_data()
    stats = calculate_stats(data)

    if stats:
        st.download_button(
            label="📥 Télécharger le Rapport (.docx)",
            data=generate_docx(stats),
            file_name="Rapport_Annuel_CMA_NKOMO.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.warning("Aucune donnée disponible.")