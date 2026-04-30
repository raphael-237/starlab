# pages/statistiques_page.py
import streamlit as st
import pandas as pd
from io import BytesIO
from modules.data_manager import load_data
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

def get_current_decade_info():
    data = load_data()
    today = datetime.now().strftime("%Y-%m-%d")
    for decade in data["decades"]:
        if decade["date_debut"] <= today <= decade["date_fin"]:
            return {
                "periode": decade["periode"],
                "date_debut": decade["date_debut"],
                "date_fin": decade["date_fin"],
                "mois": decade["mois"],
                "annee": decade["annee"]
            }
    return None

def show():
    st.markdown("""
    <style>
        .card-stat { background: white; padding: 24px; border-radius: 16px; border: 1px solid #e2e8f0; margin-bottom: 20px; }
        .stat-label { font-size: 0.8rem; color: #64748b; font-weight: 700; text-transform: uppercase; }
        .stat-val { font-size: 1.8rem; font-weight: 800; color: #0f172a; display: block; margin-top: 5px; }
        .export-section { background: #f8fafc; padding: 24px; border-radius: 16px; border: 1px dashed #cbd5e1; text-align: center; }
    </style>
    """, unsafe_allow_html=True)

    st.markdown('<div class="card p-4 border-0 shadow-sm mb-4" style="border-radius: 16px;">', unsafe_allow_html=True)
    c1, c2 = st.columns([3, 1])
    with c1:
        st.title("📊 Statistiques Financières")
        st.caption("Analysez les performances globales et générez les rapports institutionnels.")
    with c2:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🔄 Actualiser les données", use_container_width=True):
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

    data = load_data()
    if not data["transactions"]:
        st.info("ℹ️ Aucune transaction n'a encore été enregistrée.")
        return

    ex_dict = {e["id"]: e for e in data["examens"]}
    cat_dict = {c["id"]: {"nom": c["nom"], "id": c["id"]} for c in data["categories"]}
    dec_dict = {d["id"]: d for d in data["decades"]}

    records = []
    for t in data["transactions"]:
        ex = ex_dict.get(t["examen_id"])
        if not ex: continue
        nom_examen = ex["nom"]
        cat_info = cat_dict.get(ex.get("categorie_id"), {"nom": "Autres", "id": 999})
        nom_categorie = cat_info["nom"]
        categorie_id = cat_info["id"]
        if "kit" in nom_categorie.lower() or "kit" in nom_examen.lower():
            continue
        dec = dec_dict.get(t["decade_id"], {})
        montant = t.get("montant", t["quantite_payee"] * ex.get("prix", 0))
        prix_unitaire = t.get("prix_unitaire", ex.get("prix", 0))
        records.append({
            "Categorie_ID": categorie_id,
            "Catégorie": nom_categorie,
            "NOM DE L'EXAMEN": nom_examen,
            "Décade": dec.get("periode", ""),
            "Mois": dec.get("mois", ""),
            "Année": dec.get("annee", ""),
            "Quantité Payée": t["quantite_payee"],
            "Quantité Gratuite": t["quantite_gratuite"],
            "Prix Unitaire (FCFA)": prix_unitaire,
            "Discount (%)": t.get("discount_percent", 0),
            "Montant (FCFA)": montant
        })

    df = pd.DataFrame(records)
    if df.empty:
        st.warning("⚠️ Aucune donnée disponible.")
        return

    st.markdown('<div class="card p-4 border-0 shadow-sm mb-4" style="border-radius: 16px;">', unsafe_allow_html=True)
    f1, f2, f3 = st.columns([1, 1, 1])
    with f1:
        annee_sel = st.selectbox("📅 Année fiscale", sorted(df["Année"].unique(), reverse=True))
    with f2:
        mois_sel = st.selectbox("📅 Mois d'analyse", sorted(df[df["Année"] == annee_sel]["Mois"].unique()))
    with f3:
        st.markdown("<br>", unsafe_allow_html=True)
        st.success(f"Période sélectionnée : **{mois_sel} {annee_sel}**")
    st.markdown('</div>', unsafe_allow_html=True)

    df_filtered = df[(df["Mois"] == mois_sel) & (df["Année"] == annee_sel)]
    if df_filtered.empty:
        st.info(f"Aucune donnée pour {mois_sel} {annee_sel}.")
        return

    total_m = df_filtered["Montant (FCFA)"].sum()
    total_p = df_filtered["Quantité Payée"].sum()
    total_g = df_filtered["Quantité Gratuite"].sum()

    k1, k2, k3 = st.columns(3)
    with k1:
        st.markdown(f'<div class="card-stat"><span class="stat-label">Chiffre d\'Affaires</span><span class="stat-val" style="color:#0d6efd;">{total_m:,.0f} FCFA</span></div>', unsafe_allow_html=True)
    with k2:
        st.markdown(f'<div class="card-stat"><span class="stat-label">Examens Payés</span><span class="stat-val">{int(total_p)}</span></div>', unsafe_allow_html=True)
    with k3:
        st.markdown(f'<div class="card-stat"><span class="stat-label">Examens Gratuits</span><span class="stat-val">{int(total_g)}</span></div>', unsafe_allow_html=True)

    st.markdown('<div class="card p-4 border-0 shadow-sm" style="border-radius: 16px;">', unsafe_allow_html=True)
    col_table, col_exports = st.columns([2, 1], gap="large")
    with col_table:
        st.subheader("📋 Récapitulatif détaillé")
        df_display = df_filtered[["Catégorie", "NOM DE L'EXAMEN", "Décade", "Quantité Payée", "Quantité Gratuite", "Prix Unitaire (FCFA)", "Discount (%)", "Montant (FCFA)"]]
        st.dataframe(df_display, use_container_width=True, hide_index=True)
    with col_exports:
        st.subheader("📎 Rapport Officiel")
        st.markdown('<div class="export-section"><p class="text-muted small">Générez le document Word au format institutionnel.</p>', unsafe_allow_html=True)
        if st.button("📄 Générer le Rapport Word", type="primary", use_container_width=True):
            export_to_word(df_filtered, mois_sel, annee_sel, data)
            st.toast("✅ Rapport généré avec succès !", icon="📝")
        st.markdown('</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

def export_to_word(df, mois_sel, annee_sel, data):
    """
    Génère le rapport Word. Pour la catégorie Biochimie, seuls les examens
    'combi11', 'combi2', 'Glycémie' ont des valeurs ; les autres examens
    de biochimie n'affichent que leur nom sans quantités ni montants.
    """
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Inches(11.69), Inches(8.27)
    section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(0.4)

    # En-tête bilingue
    header_table = doc.add_table(rows=1, cols=3)
    header_table.width = Inches(10.5)
    def fill_header(cell, lines):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in lines:
            run = p.add_run(line + "\n")
            run.bold = True
            run.font.size = Pt(7)
    fill_header(header_table.cell(0, 0), [
        "REPUBLIQUE DU CAMEROUN", "Paix – Travail – Patrie", "------------",
        "MINISTERE DE LA SANTE PUBLIQUE", "------------",
        "SECRETARIAT GENERAL", "------------",
        "DELEGATION REGIONALE DE LA SANTE", "PUBLIQUE DU CENTRE", "------------",
        "DISTRICT DE SANTE D'ODZA", "------------",
        "CENTRE MEDICAL D'ARRONDISSEMENT DE NKOMO", "TEL. : 671742644"
    ])
    c_logo = header_table.cell(0, 1)
    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists("logo.png"):
        p_logo.add_run().add_picture("logo.png", width=Inches(1.2))
    fill_header(header_table.cell(0, 2), [
        "REPUBLIC OF CAMEROON", "Peace-Work-Fatherland", "------------",
        "MINISTRY OF PUBLIC HEALTH", "------------",
        "SECRETARIAT GENERAL", "------------",
        "CENTER REGIONAL DELEGATION OF", "PUBLIC HEALTH", "-----------",
        "ODZA HEALTH DISTRICT", "-----------",
        "NKOMO SUB-DIVISIONAL MEDICAL CENTER"
    ])

    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = t.add_run("STATISTIQUE FINANCIER DU LABORATOIRE DU CMA DE NKOMO")
    run_t.bold = True
    run_t.underline = True
    run_t.font.size = Pt(13)

    current_decade = get_current_decade_info()
    decade_text = f" - {current_decade['periode']}" if current_decade else ""
    st_title = doc.add_paragraph()
    st_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_st = st_title.add_run(f"(MOIS DE {mois_sel.upper()} {annee_sel}{decade_text})")
    run_st.bold = True

    # Préparation des données
    dec_order = ["1ère décade", "2ème décade", "3ème décade"]
    dec_labels = ["1ère décade\n(1 - 10)", "2ème décade\n(11 - 20)", "3ème décade\n(21 - 30/31)"]

    # Définir les examens de biochimie qui doivent avoir des valeurs
    biochimie_examens_avec_valeurs = ["combi11", "combi2", "Glycémie"]
    # Normaliser pour comparaison
    biochimie_examens_norm = [e.lower().strip() for e in biochimie_examens_avec_valeurs]

    # Regrouper par catégorie, examen, décade
    grouped = df.groupby(["Catégorie", "NOM DE L'EXAMEN", "Décade"]).agg({
        "Quantité Payée": "sum",
        "Quantité Gratuite": "sum",
        "Montant (FCFA)": "sum"
    }).reset_index()

    # Construire un dictionnaire pour accès facile
    pivot_data = {}
    for _, row in grouped.iterrows():
        cat = row["Catégorie"]
        exam = row["NOM DE L'EXAMEN"]
        dec = row["Décade"]
        pivot_data.setdefault(cat, {}).setdefault(exam, {})[dec] = {
            "paye": int(row["Quantité Payée"]),
            "gratuit": int(row["Quantité Gratuite"]),
            "montant": int(row["Montant (FCFA)"])
        }

    # Récupérer l'ordre des catégories
    categories_from_db = data.get("categories", [])
    category_order = {cat["id"]: idx for idx, cat in enumerate(categories_from_db)}
    unique_cats = df[["Categorie_ID", "Catégorie"]].drop_duplicates()
    unique_cats["order"] = unique_cats["Categorie_ID"].map(category_order).fillna(999)
    unique_cats = unique_cats.sort_values("order")
    categories = unique_cats["Catégorie"].tolist()
    cat_ids = unique_cats["Categorie_ID"].tolist()

    # Construction du tableau Word
    table = doc.add_table(rows=2, cols=13)
    table.style = 'Table Grid'
    h = table.rows[0].cells
    h[0].text = "NOM DE L'EXAMEN"
    for i, lbl in enumerate(dec_labels):
        h[1 + i*3].merge(h[3 + i*3]).text = lbl
    h[10].merge(h[12]).text = "TOTAL"

    sh = table.rows[1].cells
    sub_labels = ["Payé", "Gratuit", "Montant"] * 4
    for i, lbl in enumerate(sub_labels):
        sh[i+1].text = lbl

    for row in table.rows[:2]:
        for cell in row.cells:
            if cell.text:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.runs[0]
                run.bold = True
                run.font.size = Pt(8)

    grand_totals = [0] * 12

    for cat_id, cat in zip(cat_ids, categories):
        row_cat = table.add_row().cells
        cat_header = f"{cat_id}. {cat.upper()}"
        row_cat[0].text = cat_header
        row_cat[0].merge(row_cat[12])
        p_cat = row_cat[0].paragraphs[0]
        p_cat.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cat = p_cat.runs[0]
        run_cat.bold = True
        run_cat.underline = True
        run_cat.font.size = Pt(9)

        examens = sorted(pivot_data.get(cat, {}).keys())
        cat_sums = [0] * 9

        for exam in examens:
            rc = table.add_row().cells
            rc[0].text = exam
            ex_paye, ex_gratuit, ex_montant = 0, 0, 0

            # Déterminer si l'examen doit avoir des valeurs
            is_biochimie = (cat.lower() == "biochimie")
            est_examen_avec_valeurs = (exam.lower().strip() in biochimie_examens_norm)

            for idx_dec, dec in enumerate(dec_order):
                col_start = 1 + idx_dec * 3
                if is_biochimie and not est_examen_avec_valeurs:
                    # Pour les examens de biochimie non-listés, on laisse les cellules vides
                    rc[col_start].text = ""
                    rc[col_start+1].text = ""
                    rc[col_start+2].text = ""
                else:
                    vals = pivot_data[cat][exam].get(dec, {"paye":0, "gratuit":0, "montant":0})
                    rc[col_start].text = str(vals["paye"]) if vals["paye"] else ""
                    rc[col_start+1].text = str(vals["gratuit"]) if vals["gratuit"] else ""
                    rc[col_start+2].text = f"{vals['montant']:,}" if vals["montant"] else ""
                    ex_paye += vals["paye"]
                    ex_gratuit += vals["gratuit"]
                    ex_montant += vals["montant"]
                    cat_sums[idx_dec*3] += vals["paye"]
                    cat_sums[idx_dec*3+1] += vals["gratuit"]
                    cat_sums[idx_dec*3+2] += vals["montant"]

            # Colonnes TOTAL
            if is_biochimie and not est_examen_avec_valeurs:
                rc[10].text = ""
                rc[11].text = ""
                rc[12].text = ""
            else:
                rc[10].text = str(ex_paye) if ex_paye else ""
                rc[11].text = str(ex_gratuit) if ex_gratuit else ""
                rc[12].text = f"{ex_montant:,}" if ex_montant else ""

            for cell in rc:
                p = cell.paragraphs[0]
                if p.runs:
                    p.runs[0].font.size = Pt(8)
                if cell != rc[0]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Sous-total de la catégorie (seulement des examens avec valeurs contribuent)
        rst = table.add_row().cells
        rst[0].text = f"Sous-Total {cat_header}"
        total_paye = cat_sums[0] + cat_sums[3] + cat_sums[6]
        total_gratuit = cat_sums[1] + cat_sums[4] + cat_sums[7]
        total_montant = cat_sums[2] + cat_sums[5] + cat_sums[8]
        all_vals = cat_sums + [total_paye, total_gratuit, total_montant]
        for i in range(12):
            val = all_vals[i]
            if i % 3 == 2:
                rst[1+i].text = f"{int(val):,}" if val else ""
            else:
                rst[1+i].text = str(int(val)) if val else ""
            grand_totals[i] += val
        for cell in rst:
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if not p.runs:
                p.add_run(cell.text)
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(8)

    # Total général
    rtot = table.add_row().cells
    rtot[0].text = "TOTAL GÉNÉRAL"
    for i in range(12):
        val = grand_totals[i]
        if i % 3 == 2:
            rtot[1+i].text = f"{int(val):,}" if val else ""
        else:
            rtot[1+i].text = str(int(val)) if val else ""
        p = rtot[1+i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not p.runs:
            p.add_run(rtot[1+i].text)
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(8)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    st.download_button("⬇️ Télécharger le Rapport Word Officiel", output,
                       file_name=f"RAPPORT_LABO_{mois_sel.upper()}_{annee_sel}.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")